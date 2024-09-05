from docx import Document
from docx.shared import Pt

def add_section(doc, title, content):
    """Add a section with a title and content to the document."""
    doc.add_heading(title, level=1)
    doc.add_paragraph(content)
    doc.add_paragraph()  

def create_social_media_policy_document(app_name, terms, privacy_policy, acceptable_use_policy, user_rights, contact_info):
    """Create a comprehensive policy document for a social media app."""
    
    doc = Document()
    doc.add_heading(f'{app_name} Policy Document', 0)

    add_section(doc, 'Terms of Service', terms)

    add_section(doc, 'Privacy Policy', privacy_policy)
    
    add_section(doc, 'Acceptable Use Policy', acceptable_use_policy)
    
    add_section(doc, 'User Rights', user_rights)
    
    add_section(doc, 'Contact Information', contact_info)
  
    filename = f'{app_name}_Policy_Document.docx'
    doc.save(filename)
    print(f'Policy document "{filename}" has been created.')

def gather_policy_details():
    """Gather policy details from the user."""
    app_name = input('Enter the name of the social media app: ')

    print("\nEnter the Terms of Service:")
    terms = input()

    print("\nEnter the Privacy Policy:")
    privacy_policy = input()

    print("\nEnter the Acceptable Use Policy:")
    acceptable_use_policy = input()

    print("\nEnter the User Rights:")
    user_rights = input()

    print("\nEnter the Contact Information:")
    contact_info = input()

    return app_name, terms, privacy_policy, acceptable_use_policy, user_rights, contact_info

def main():
    print("Welcome to the Social Media Policy Document Generator")

    app_name, terms, privacy_policy, acceptable_use_policy, user_rights, contact_info = gather_policy_details()
    create_social_media_policy_document(app_name, terms, privacy_policy, acceptable_use_policy, user_rights, contact_info)

if __name__ == "__main__":
    main()
